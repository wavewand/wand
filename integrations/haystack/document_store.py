"""
Haystack Document Store Management for MCP System

This module handles document ingestion, storage, and retrieval for the Haystack
integration, supporting multiple document formats and vector databases.
"""

import hashlib
import logging
import mimetypes
import uuid
from datetime import datetime
from typing import Any, BinaryIO, Dict, List, Optional

try:
    from haystack import Document
    from haystack.document_stores.in_memory import InMemoryDocumentStore

    HAYSTACK_AVAILABLE = True
except ImportError:
    logging.warning("Haystack not available. Install with: pip install haystack-ai")
    HAYSTACK_AVAILABLE = False

    class Document:
        def __init__(self, *args, **kwargs):
            pass

    class InMemoryDocumentStore:
        def __init__(self, *args, **kwargs):
            pass


# Document processing imports
try:
    import markdown
    import pypdf
    from bs4 import BeautifulSoup
    from docx import Document as DocxDocument

    PDF_AVAILABLE = True
    DOCX_AVAILABLE = True
    HTML_AVAILABLE = True
    MARKDOWN_AVAILABLE = True
except ImportError as e:
    logging.warning(f"Some document processing libraries not available: {e}")
    PDF_AVAILABLE = False
    DOCX_AVAILABLE = False
    HTML_AVAILABLE = False
    MARKDOWN_AVAILABLE = False


class HaystackDocumentStore:
    """Manages document storage and processing for Haystack integration."""

    def __init__(self, vector_db_config: Optional[Dict[str, Any]] = None):
        self.logger = logging.getLogger(__name__)
        self.document_store = None
        self.vector_db_config = vector_db_config or {}
        self.processed_documents: Dict[str, Dict[str, Any]] = {}

        self._initialize_document_store()

    def _initialize_document_store(self):
        """Initialize the document store based on configuration."""
        if not HAYSTACK_AVAILABLE:
            self.logger.warning("Haystack not available - document store disabled")
            return

        try:
            # For now, use InMemoryDocumentStore
            # Future: Add support for vector databases like Weaviate, Pinecone, Qdrant
            self.document_store = InMemoryDocumentStore()
            self.logger.info("Initialized in-memory document store")

        except Exception as e:
            self.logger.error(f"Failed to initialize document store: {e}")

    def ingest_document(
        self, filename: str, content: bytes, content_type: str = None, metadata: Dict[str, Any] = None
    ) -> Dict[str, Any]:
        """Ingest a document into the document store."""
        if not self.document_store:
            return {"success": False, "message": "Document store not available"}

        try:
            # Determine content type if not provided
            if not content_type:
                content_type, _ = mimetypes.guess_type(filename)
                content_type = content_type or "text/plain"

            # Process document based on type
            processed_content = self._process_document_content(content, content_type, filename)

            if not processed_content:
                return {"success": False, "message": f"Failed to process document of type {content_type}"}

            # Create document ID and metadata
            doc_id = str(uuid.uuid4())
            doc_metadata = {
                "filename": filename,
                "content_type": content_type,
                "ingested_at": datetime.now().isoformat(),
                "size_bytes": len(content),
                "content_hash": hashlib.sha256(content).hexdigest(),
                **(metadata or {}),
            }

            # Create Haystack document
            document = Document(content=processed_content, meta=doc_metadata, id=doc_id)

            # Store in document store
            self.document_store.write_documents([document])

            # Track processed document
            self.processed_documents[doc_id] = {
                "id": doc_id,
                "filename": filename,
                "content_type": content_type,
                "metadata": doc_metadata,
                "created_at": datetime.now(),
            }

            self.logger.info(f"Successfully ingested document: {filename}")

            return {
                "success": True,
                "document_id": doc_id,
                "message": f"Document {filename} ingested successfully",
                "content_length": len(processed_content),
                "metadata": doc_metadata,
            }

        except Exception as e:
            self.logger.error(f"Failed to ingest document {filename}: {e}")
            return {"success": False, "message": f"Failed to ingest document: {str(e)}"}

    def _process_document_content(self, content: bytes, content_type: str, filename: str) -> Optional[str]:
        """Process document content based on its type."""
        try:
            if content_type == "application/pdf":
                return self._process_pdf(content)
            elif content_type in [
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                "application/msword",
            ]:
                return self._process_docx(content)
            elif content_type in ["text/html", "application/xhtml+xml"]:
                return self._process_html(content)
            elif content_type == "text/markdown" or filename.endswith(('.md', '.markdown')):
                return self._process_markdown(content)
            elif content_type.startswith("text/"):
                return content.decode('utf-8', errors='ignore')
            else:
                # Try to decode as text
                return content.decode('utf-8', errors='ignore')

        except Exception as e:
            self.logger.error(f"Error processing {content_type}: {e}")
            return None

    def _process_pdf(self, content: bytes) -> Optional[str]:
        """Extract text from PDF content."""
        if not PDF_AVAILABLE:
            self.logger.warning("PDF processing not available - install pypdf")
            return None

        try:
            import io

            pdf_file = io.BytesIO(content)
            reader = pypdf.PdfReader(pdf_file)

            text_content = []
            for page in reader.pages:
                text_content.append(page.extract_text())

            return "\n".join(text_content)

        except Exception as e:
            self.logger.error(f"PDF processing failed: {e}")
            return None

    def _process_docx(self, content: bytes) -> Optional[str]:
        """Extract text from DOCX content."""
        if not DOCX_AVAILABLE:
            self.logger.warning("DOCX processing not available - install python-docx")
            return None

        try:
            import io

            docx_file = io.BytesIO(content)
            doc = DocxDocument(docx_file)

            text_content = []
            for paragraph in doc.paragraphs:
                text_content.append(paragraph.text)

            return "\n".join(text_content)

        except Exception as e:
            self.logger.error(f"DOCX processing failed: {e}")
            return None

    def _process_html(self, content: bytes) -> Optional[str]:
        """Extract text from HTML content."""
        if not HTML_AVAILABLE:
            self.logger.warning("HTML processing not available - install beautifulsoup4")
            return None

        try:
            html_text = content.decode('utf-8', errors='ignore')
            soup = BeautifulSoup(html_text, 'html.parser')

            # Remove script and style elements
            for script in soup(["script", "style"]):
                script.decompose()

            return soup.get_text(separator='\n', strip=True)

        except Exception as e:
            self.logger.error(f"HTML processing failed: {e}")
            return None

    def _process_markdown(self, content: bytes) -> Optional[str]:
        """Process Markdown content."""
        if not MARKDOWN_AVAILABLE:
            self.logger.warning("Markdown processing not available - install markdown")
            return content.decode('utf-8', errors='ignore')

        try:
            md_text = content.decode('utf-8', errors='ignore')
            html = markdown.markdown(md_text)

            # Convert HTML to plain text
            if HTML_AVAILABLE:
                soup = BeautifulSoup(html, 'html.parser')
                return soup.get_text(separator='\n', strip=True)
            else:
                return md_text

        except Exception as e:
            self.logger.error(f"Markdown processing failed: {e}")
            return content.decode('utf-8', errors='ignore')

    def get_document(self, document_id: str) -> Optional[Dict[str, Any]]:
        """Retrieve a document by ID."""
        if not self.document_store:
            return None

        try:
            documents = self.document_store.filter_documents(filters={"id": document_id})
            if documents:
                doc = documents[0]
                return {"id": doc.id, "content": doc.content, "metadata": doc.meta}
            return None

        except Exception as e:
            self.logger.error(f"Failed to get document {document_id}: {e}")
            return None

    def list_documents(self, filters: Dict[str, Any] = None, limit: int = 100) -> List[Dict[str, Any]]:
        """List documents with optional filters."""
        if not self.document_store:
            return []

        try:
            documents = self.document_store.filter_documents(filters=filters or {})

            results = []
            for doc in documents[:limit]:
                results.append(
                    {
                        "id": doc.id,
                        "filename": doc.meta.get("filename", "unknown"),
                        "content_type": doc.meta.get("content_type", "text/plain"),
                        "ingested_at": doc.meta.get("ingested_at"),
                        "size_bytes": doc.meta.get("size_bytes", 0),
                        "content_preview": doc.content[:200] + "..." if len(doc.content) > 200 else doc.content,
                    }
                )

            return results

        except Exception as e:
            self.logger.error(f"Failed to list documents: {e}")
            return []

    def delete_document(self, document_id: str) -> Dict[str, Any]:
        """Delete a document from the store."""
        if not self.document_store:
            return {"success": False, "message": "Document store not available"}

        try:
            # Check if document exists
            documents = self.document_store.filter_documents(filters={"id": document_id})
            if not documents:
                return {"success": False, "message": f"Document {document_id} not found"}

            # Delete from document store
            self.document_store.delete_documents([document_id])

            # Remove from tracking
            if document_id in self.processed_documents:
                del self.processed_documents[document_id]

            return {"success": True, "message": f"Document {document_id} deleted successfully"}

        except Exception as e:
            self.logger.error(f"Failed to delete document {document_id}: {e}")
            return {"success": False, "message": f"Failed to delete document: {str(e)}"}

    def update_document_metadata(self, document_id: str, metadata: Dict[str, Any]) -> Dict[str, Any]:
        """Update document metadata."""
        if not self.document_store:
            return {"success": False, "message": "Document store not available"}

        try:
            # Get existing document
            documents = self.document_store.filter_documents(filters={"id": document_id})
            if not documents:
                return {"success": False, "message": f"Document {document_id} not found"}

            doc = documents[0]

            # Update metadata
            updated_meta = {**doc.meta, **metadata}
            updated_doc = Document(content=doc.content, meta=updated_meta, id=doc.id)

            # Replace document
            self.document_store.delete_documents([document_id])
            self.document_store.write_documents([updated_doc])

            return {"success": True, "message": f"Document {document_id} metadata updated", "metadata": updated_meta}

        except Exception as e:
            self.logger.error(f"Failed to update document metadata {document_id}: {e}")
            return {"success": False, "message": f"Failed to update metadata: {str(e)}"}

    def get_stats(self) -> Dict[str, Any]:
        """Get document store statistics."""
        if not self.document_store:
            return {"total_documents": 0, "document_store_available": False}

        try:
            all_docs = self.document_store.filter_documents()

            content_types = {}
            total_size = 0

            for doc in all_docs:
                content_type = doc.meta.get("content_type", "unknown")
                content_types[content_type] = content_types.get(content_type, 0) + 1
                total_size += doc.meta.get("size_bytes", 0)

            return {
                "total_documents": len(all_docs),
                "content_types": content_types,
                "total_size_bytes": total_size,
                "document_store_available": True,
                "processing_capabilities": {
                    "pdf": PDF_AVAILABLE,
                    "docx": DOCX_AVAILABLE,
                    "html": HTML_AVAILABLE,
                    "markdown": MARKDOWN_AVAILABLE,
                },
            }

        except Exception as e:
            self.logger.error(f"Failed to get stats: {e}")
            return {"total_documents": 0, "document_store_available": False, "error": str(e)}
